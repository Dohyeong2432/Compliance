import io

from docx import Document

from agent.contract_docx import build_review_document
from agent.contract_review import ClauseReview
from agent.harness import AgentTurnResult
from agent.sso import SessionContext


def _paragraph_texts(document) -> list[str]:
    return [p.text for p in document.paragraphs]


def test_build_review_document_includes_title_and_metadata():
    session = SessionContext(user_id="tester1", dept="IB")
    document = build_review_document("업무위탁계약서.docx", session, [])

    texts = _paragraph_texts(document)
    assert document.paragraphs[0].text == "계약서 검토의견서" or any(
        h.text == "계약서 검토의견서" for h in document.paragraphs
    )
    assert any("업무위탁계약서.docx" in t for t in texts)
    assert any("tester1" in t and "IB" in t for t in texts)


def test_build_review_document_includes_clause_label_original_and_answer():
    session = SessionContext(user_id="u1", dept="RETAIL")
    reviews = [
        ClauseReview(
            label="제1조(목적)",
            original_text="이 계약은 위수탁 업무 범위를 정한다.",
            result=AgentTurnResult(answer="법령 위반 소지가 없습니다."),
        )
    ]

    document = build_review_document("계약서.docx", session, reviews)
    texts = _paragraph_texts(document)

    assert "제1조(목적)" in texts
    assert "이 계약은 위수탁 업무 범위를 정한다." in texts
    assert "법령 위반 소지가 없습니다." in texts


def test_build_review_document_renders_multiple_clauses_in_order():
    session = SessionContext(user_id="u1", dept="RETAIL")
    reviews = [
        ClauseReview("제1조(목적)", "원문1", AgentTurnResult(answer="검토1")),
        ClauseReview("제2조(정의)", "원문2", AgentTurnResult(answer="검토2")),
    ]

    document = build_review_document("계약서.docx", session, reviews)
    texts = _paragraph_texts(document)

    assert texts.index("제1조(목적)") < texts.index("제2조(정의)")
    assert texts.index("검토1") < texts.index("검토2")


def test_build_review_document_formats_citation_reference_list():
    """CitationGuard.apply()가 만든 "**참고 문서**" 마커는 볼드 처리되고,
    각주 참고문서 줄은 그대로 문단으로 남아야 한다."""
    session = SessionContext(user_id="u1", dept="RETAIL")
    answer = "문제 없습니다 [1].\n\n---\n**참고 문서**\n[1] 금융지주회사법 제47조 (law:47-0)"
    reviews = [ClauseReview("제1조(목적)", "원문", AgentTurnResult(answer=answer))]

    document = build_review_document("계약서.docx", session, reviews)
    texts = _paragraph_texts(document)

    assert "참고 문서" in texts
    assert "[1] 금융지주회사법 제47조 (law:47-0)" in texts
    assert "---" not in texts  # 구분선은 문서 문단으로 남기지 않음


def test_build_review_document_round_trips_through_docx_bytes():
    """실제 .docx 바이트로 저장 후 다시 읽어도 내용이 보존되는지 -- API
    엔드포인트가 정확히 이 왕복(save to BytesIO -> Response)을 한다."""
    session = SessionContext(user_id="u1", dept="RETAIL")
    reviews = [ClauseReview("제1조(목적)", "원문", AgentTurnResult(answer="검토의견"))]
    document = build_review_document("계약서.docx", session, reviews)

    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)

    reloaded = Document(buffer)
    assert "검토의견" in _paragraph_texts(reloaded)
