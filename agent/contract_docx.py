"""ClauseReview 목록을 검토의견서 .docx로 조립한다.

원본 계약서 파일에 코멘트를 삽입하는 대신 별도 문서를 새로 만든다 --
python-docx의 네이티브 코멘트 API는 저수준 OOXML 조작이 필요해 원본 문서
구조를 깨뜨릴 위험이 있고, 별도 문서는 조항/원문/검토의견/근거를 표 없이도
읽기 쉬운 순서로 정리할 수 있다.
"""

from __future__ import annotations

from datetime import datetime

from docx import Document

from agent.contract_review import ClauseReview
from agent.sso import SessionContext

_REFERENCE_LIST_MARKER = "**참고 문서**"

# agent.contract_review.CLAUSE_REVIEW_PROMPT_TEMPLATE이 요구하는 구조화 필드
# 라벨. 이 줄들만 라벨 부분을 볼드 처리해 검토의견서에서 한눈에 훑어볼 수
# 있게 한다 -- 체크리스트 도입으로 답변이 길어진 만큼, 문서에서도 구조가
# 보여야 실제로 "세밀해진" 효과가 있다.
_FIELD_LABELS = ("위험도:", "문제 조항:", "근거:", "수정 제안:")


def _add_answer_paragraphs(document: Document, answer: str) -> None:
    """agent.ask()가 돌려주는 answer는 이미 CitationGuard.apply()가
    [[CITE:id]]를 각주 번호로, 하단에 "**참고 문서**" 목록으로 정리해준
    텍스트다. 별도 마크다운 파서 없이 줄 단위로 문단화하되, 참고 문서 제목
    줄과 구조화 필드 라벨(위험도/문제 조항/근거/수정 제안)만 볼드 처리한다."""
    for line in answer.splitlines():
        stripped = line.strip()
        if not stripped:
            continue
        if stripped == _REFERENCE_LIST_MARKER:
            document.add_paragraph().add_run("참고 문서").bold = True
            continue
        if stripped == "---":
            continue

        label = next((field_label for field_label in _FIELD_LABELS if stripped.startswith(field_label)), None)
        if label is not None:
            paragraph = document.add_paragraph()
            paragraph.add_run(label).bold = True
            paragraph.add_run(stripped[len(label):])
        else:
            document.add_paragraph(line)


def build_review_document(
    source_filename: str, session: SessionContext, clause_reviews: list[ClauseReview]
) -> Document:
    document = Document()

    document.add_heading("계약서 검토의견서", level=0)
    document.add_paragraph(f"원본 파일명: {source_filename}")
    document.add_paragraph(f"검토일시: {datetime.now():%Y-%m-%d %H:%M}")
    document.add_paragraph(f"요청자: {session.user_id} ({session.dept})")

    for review in clause_reviews:
        document.add_heading(review.label, level=1)

        document.add_paragraph().add_run("원문").bold = True
        document.add_paragraph(review.original_text)

        document.add_paragraph().add_run("검토의견").bold = True
        _add_answer_paragraphs(document, review.result.answer)

    return document
